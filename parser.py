import pdfplumber
import docx
import io
import pypdfium2 as pdfium
from PIL import Image
from pdfminer.high_level import extract_text as pdfminer_extract_text

try:
    import pytesseract
except ImportError:
    pytesseract = None

import base64
import os
import re
from openai import OpenAI

MIN_USABLE_TEXT_LENGTH = 100


def normalize_extracted_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def log_extracted_text(label: str, text: str) -> None:
    print(f"\n[PDF DEBUG] {label} extracted {len(text)} characters")
    print("[PDF DEBUG] RAW_EXTRACTED_TEXT_START")
    print(text[:4000] if text else "<EMPTY>")
    print("[PDF DEBUG] RAW_EXTRACTED_TEXT_END\n")

def parse_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                print(f"[PDF DEBUG] pdfplumber page {index}: {len(extracted.strip())} characters")
                if extracted:
                    text_parts.append(extracted)
    except Exception as e:
        print(f"PDF Parse Error: {e}")

    extracted_text = normalize_extracted_text("\n".join(text_parts))
    log_extracted_text("pdfplumber", extracted_text)

    if len(extracted_text) < MIN_USABLE_TEXT_LENGTH:
        try:
            pdfminer_text = normalize_extracted_text(pdfminer_extract_text(io.BytesIO(file_bytes)) or "")
            log_extracted_text("pdfminer fallback", pdfminer_text)
            if len(pdfminer_text) > len(extracted_text):
                extracted_text = pdfminer_text
        except Exception as e:
            print(f"PDFMiner Parse Error: {e}")

    print("Extracted characters:", len(extracted_text))
    print("First 500 chars:", extracted_text[:500] if extracted_text else "<EMPTY>")
    return extracted_text

def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"DOCX Parse Error: {e}")
        return ""


def run_vision_ocr(image: Image.Image) -> str:
    # Groq doesn't support vision — only use if OpenAI or Gemini key available
    openai_key = os.environ.get("OPENAI_API_KEY")
    gemini_key = os.environ.get("GEMINI_API_KEY")
    api_key = openai_key or gemini_key
    if not api_key:
        return ""
    try:
        if gemini_key:
            from openai import OpenAI as _OAI
            _client = _OAI(api_key=gemini_key, base_url="https://generativelanguage.googleapis.com/v1beta/openai/")
            model = "gemini-2.0-flash"
        else:
            from openai import OpenAI as _OAI
            _client = _OAI(api_key=openai_key)
            model = "gpt-4o-mini"
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        base64_image = base64.b64encode(buffered.getvalue()).decode('utf-8')
        response = _client.chat.completions.create(
            model=model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "Extract all text from this resume exactly as it appears. Maintain structure. Return only the extracted text."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]
            }],
            max_tokens=2000
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        print(f"Vision OCR Error: {e}")
        return ""

def run_ocr_on_image(image: Image.Image) -> dict:
    text = ""
    error = None
    ocr_available = pytesseract is not None
    if pytesseract is not None:
        try:
            text = pytesseract.image_to_string(image)
        except Exception as e:
            error = str(e)
    if not text.strip():
        vision_text = run_vision_ocr(image)
        if vision_text.strip():
            text = vision_text
            error = None
    return {
        "text": normalize_extracted_text(text),
        "ocr_available": ocr_available,
        "error": error if not text.strip() else None,
    }


def parse_image_with_ocr(file_bytes: bytes) -> dict:
    try:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    except Exception as e:
        return {
            "text": "",
            "ocr_available": pytesseract is not None,
            "error": f"Image Parse Error: {e}",
        }

    return run_ocr_on_image(image)


def parse_scanned_pdf_with_ocr(file_bytes: bytes, max_pages: int = 3) -> dict:

    try:
        pdf = pdfium.PdfDocument(file_bytes)
        text_parts = []
        page_count = min(len(pdf), max_pages)
        for page_index in range(page_count):
            image = pdf[page_index].render(scale=1.5).to_pil().convert("RGB")
            ocr_result = run_ocr_on_image(image)
            if ocr_result["error"]:
                return ocr_result
            if ocr_result["text"].strip():
                text_parts.append(ocr_result["text"])

        extracted_text = normalize_extracted_text("\n".join(text_parts))
        log_extracted_text("OCR fallback", extracted_text)
        return {
            "text": extracted_text,
            "ocr_available": True,
            "error": None,
        }
    except Exception as e:
        return {
            "text": "",
            "ocr_available": True,
            "error": f"PDF OCR Error: {e}",
        }


def analyze_image_layout(file_bytes: bytes) -> dict:
    try:
        image = Image.open(io.BytesIO(file_bytes)).convert("L")
        width, height = image.size
        total_pixels = max(width * height, 1)
        histogram = image.histogram()
        dark_ratio = sum(histogram[:180]) / total_pixels
        white_ratio = sum(histogram[241:]) / total_pixels

        row_dark_ratios = []
        pixels = image.load()
        for y in range(height):
            dark_pixels = 0
            for x in range(width):
                if pixels[x, y] < 180:
                    dark_pixels += 1
            row_dark_ratios.append(dark_pixels / width)

        text_band_count = 0
        in_band = False
        for ratio in row_dark_ratios:
            if ratio > 0.05 and not in_band:
                text_band_count += 1
                in_band = True
            elif ratio <= 0.05 and in_band:
                in_band = False

        horizontal_rule_rows = sum(1 for ratio in row_dark_ratios if ratio > 0.45)

        is_resume_like_image = (
            0.03 <= dark_ratio <= 0.22
            and white_ratio >= 0.55
            and text_band_count >= 12
            and horizontal_rule_rows >= 1
        )

        return {
            "is_resume_like_image": is_resume_like_image,
            "width": width,
            "height": height,
            "dark_ratio": round(dark_ratio, 3),
            "white_ratio": round(white_ratio, 3),
            "text_band_count": text_band_count,
            "horizontal_rule_rows": horizontal_rule_rows,
        }
    except Exception as e:
        print(f"Image Layout Analysis Error: {e}")
        return {"is_resume_like_image": False, "error": str(e)}


def analyze_scanned_pdf_layout(file_bytes: bytes) -> dict:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            if page_count == 0:
                return {"is_scanned_resume": False}

            first_page = pdf.pages[0]
            page_area = max(first_page.width * first_page.height, 1)
            largest_image_coverage = max(
                ((image.get("width", 0) * image.get("height", 0)) / page_area for image in first_page.images),
                default=0,
            )
            image_count = len(first_page.images)
            char_count = len(first_page.chars)

        pdf = pdfium.PdfDocument(file_bytes)
        rendered = pdf[0].render(scale=1.5).to_pil().convert("L")
        width, height = rendered.size
        histogram = rendered.histogram()
        total_pixels = max(width * height, 1)
        dark_ratio = sum(histogram[:180]) / total_pixels
        white_ratio = sum(histogram[241:]) / total_pixels

        row_dark_ratios = []
        pixels = rendered.load()
        for y in range(height):
            dark_pixels = 0
            for x in range(width):
                if pixels[x, y] < 180:
                    dark_pixels += 1
            row_dark_ratios.append(dark_pixels / width)

        text_band_count = 0
        in_band = False
        for ratio in row_dark_ratios:
            if ratio > 0.05 and not in_band:
                text_band_count += 1
                in_band = True
            elif ratio <= 0.05 and in_band:
                in_band = False

        horizontal_rule_rows = sum(1 for ratio in row_dark_ratios if ratio > 0.45)

        is_resume_like_scan = (
            page_count <= 2
            and char_count == 0
            and image_count >= 1
            and largest_image_coverage > 0.75
            and 0.03 <= dark_ratio <= 0.18
            and white_ratio >= 0.70
            and text_band_count >= 18
            and horizontal_rule_rows >= 4
        )

        return {
            "is_scanned_resume": is_resume_like_scan,
            "page_count": page_count,
            "image_count": image_count,
            "char_count": char_count,
            "largest_image_coverage": round(largest_image_coverage, 2),
            "dark_ratio": round(dark_ratio, 3),
            "white_ratio": round(white_ratio, 3),
            "text_band_count": text_band_count,
            "horizontal_rule_rows": horizontal_rule_rows,
        }
    except Exception as e:
        print(f"Scan Layout Analysis Error: {e}")
        return {"is_scanned_resume": False, "error": str(e)}
